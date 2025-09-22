from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from models import ExportRequest
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import black
import io
import os

router = APIRouter()

def create_download_file(content: str, filename: str, file_type: str) -> bytes:
    """Create downloadable file content"""
    if file_type == "docx":
        # Create a well-formatted DOCX file
        doc = Document()
        
        # Set default font
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.shared import OxmlElement, qn
        
        # Process content with better formatting
        lines = content.split('\n')
        i = 0
        in_skills_section = False
        
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            
            # Check if it's a name (first non-empty line, all caps)
            if i == 0 or (line.isupper() and len(line) > 3 and not any(char in line for char in ['|', '@', '.', ':', '•'])):
                name_para = doc.add_paragraph()
                name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                name_run = name_para.add_run(line)
                # 14px for resume, 12px for cover letter (matching popup)
                name_font_size = 14 if 'resume' in filename.lower() else 12
                name_run.font.size = Pt(name_font_size)
                name_run.font.bold = True
                name_run.font.name = 'Arial'
                doc.add_paragraph()  # Add space after name (mb-3 equivalent)
            
            # Check if it's contact info
            elif ('@' in line or 'linkedin.com' in line.lower() or 
                  (any(char.isdigit() for char in line) and len(line) > 10) or
                  'github.com' in line.lower()):
                contact_para = doc.add_paragraph()
                contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                contact_run = contact_para.add_run(line)
                contact_run.font.size = Pt(11)  # text-sm equivalent
                contact_run.font.name = 'Arial'
                doc.add_paragraph()  # Add space after contact
            
            # Check if it's a section header
            elif (line.isupper() and len(line) > 3 and 
                  any(section in line for section in [
                      'PROFESSIONAL SUMMARY', 'WORK EXPERIENCE', 'EDUCATION', 
                      'SKILLS', 'PROJECTS', 'CERTIFICATIONS', 'EXTRACURRICULAR ACTIVITIES',
                      'AWARDS & ACHIEVEMENTS', 'VOLUNTEER', 'LANGUAGES', 'INTERESTS',
                      'TECHNICAL SKILLS', 'PROFESSIONAL EXPERIENCE'
                  ])):
                in_skills_section = 'SKILLS' in line
                # Add margin-top: 6px before section
                doc.add_paragraph()  # This adds the 6px margin-top
                
                section_para = doc.add_paragraph()
                section_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                section_run = section_para.add_run(line)
                # 12px for resume, 12px for cover letter (matching popup)
                section_font_size = 12 if 'resume' in filename.lower() else 12
                section_run.font.size = Pt(section_font_size)
                section_run.font.bold = True
                section_run.font.name = 'Arial'
                
                # Add line under heading (like in popup)
                line_para = doc.add_paragraph()
                line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                line_run = line_para.add_run('_' * 50)  # Add line under heading
                line_run.font.size = Pt(8)
                line_run.font.name = 'Arial'
                
                # Add 6px gap after heading (increased by 2px)
                doc.add_paragraph()  # This adds the 6px gap after heading
            
            # Check if it's a job title/company line
            elif '|' in line and len(line.split('|')) >= 2:
                job_para = doc.add_paragraph()
                job_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                job_run = job_para.add_run(line)
                # 10px for resume, 10px for cover letter
                job_font_size = 10 if 'resume' in filename.lower() else 10
                job_run.font.size = Pt(job_font_size)
                job_run.font.bold = True
                job_run.font.name = 'Arial'
                # Add 6px gap after job title (increased by 2px)
                doc.add_paragraph()  # This adds the 6px gap after job title
            
            # Check if it's a bullet point
            elif line.startswith('•') or line.startswith('-'):
                bullet_para = doc.add_paragraph()
                bullet_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                bullet_run = bullet_para.add_run(line)  # Keep bullet character
                # 9px for resume, 10px for cover letter
                bullet_font_size = 9 if 'resume' in filename.lower() else 10
                bullet_run.font.size = Pt(bullet_font_size)
                bullet_run.font.name = 'Arial'
                # Add 6px gap after bullet point (increased by 2px)
                doc.add_paragraph()  # This adds the 6px gap after bullet point
            
            # Check if it's in skills section
            elif in_skills_section and ':' in line:
                skills_para = doc.add_paragraph()
                skills_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                skills_run = skills_para.add_run(line)
                # 9px for resume, 10px for cover letter
                skills_font_size = 9 if 'resume' in filename.lower() else 10
                skills_run.font.size = Pt(skills_font_size)
                skills_run.font.name = 'Arial'
                # Add 6px gap after skills line (increased by 2px)
                doc.add_paragraph()  # This adds the 6px gap after skills line
            
            # Regular text
            else:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = para.add_run(line)
                # 9px for resume, 10px for cover letter
                normal_font_size = 9 if 'resume' in filename.lower() else 10
                run.font.size = Pt(normal_font_size)
                run.font.name = 'Arial'
                # Add 6px gap after regular text (increased by 2px)
                doc.add_paragraph()  # This adds the 6px gap after regular text
            
            i += 1
        
        # Save to bytes
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io.getvalue()
    elif file_type == "pdf":
        try:
            # Create PDF in memory with better margins
            pdf_io = io.BytesIO()
            doc = SimpleDocTemplate(
                pdf_io, 
                pagesize=letter,
                leftMargin=0.75*inch, 
                rightMargin=0.75*inch,
                topMargin=0.75*inch, 
                bottomMargin=0.75*inch
            )
            
            # Create custom styles with consistent, professional formatting
            styles = getSampleStyleSheet()
            
            # CONSISTENT FONT SIZE SYSTEM
            # Name: 16px (largest, most prominent)
            # Section Headers: 12px (clear hierarchy)
            # Job Titles: 11px (important but not overwhelming)
            # Contact Info: 10px (supporting information)
            # Body Text: 10px (main content)
            # Skills: 10px (consistent with body)
            
            # Name style - 16px for both resume and cover letter (consistent, prominent)
            name_style = ParagraphStyle(
                'NameStyle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=8,
                spaceBefore=0,
                alignment=0,  # Left align
                textColor=black,
                fontName='Helvetica-Bold',
                leading=18
            )
            
            # Contact style - 10px (consistent with body text)
            contact_style = ParagraphStyle(
                'ContactStyle',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=12,
                spaceBefore=0,
                alignment=0,  # Left align
                fontName='Helvetica',
                leading=12
            )
            
            # Section header style - 12px (clear hierarchy, consistent)
            section_style = ParagraphStyle(
                'SectionStyle',
                parent=styles['Heading2'],
                fontSize=12,
                spaceAfter=4,
                spaceBefore=20,
                textColor=black,
                fontName='Helvetica-Bold',
                alignment=0,  # Left align
                leftIndent=0,
                rightIndent=0,
                leading=14
            )
            
            # Job title/company style - 11px (important but not overwhelming)
            job_title_style = ParagraphStyle(
                'JobTitleStyle',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=4,
                spaceBefore=0,
                fontName='Helvetica-Bold',
                leftIndent=0,
                rightIndent=0,
                alignment=0,  # Left align
                leading=13,
                textColor=black
            )
            
            # Normal text style - 10px (main content, consistent)
            normal_style = ParagraphStyle(
                'NormalStyle',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=4,
                spaceBefore=0,
                fontName='Helvetica',
                leftIndent=0,
                rightIndent=0,
                alignment=0,  # Left align
                leading=12
            )
            
            # Bullet point style - 10px (consistent with body text)
            bullet_style = ParagraphStyle(
                'BulletStyle',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=4,
                spaceBefore=0,
                fontName='Helvetica',
                leftIndent=18,
                rightIndent=0,
                alignment=0,  # Left align
                leading=12
            )
            
            # Skills style - 10px (consistent with body text)
            skills_style = ParagraphStyle(
                'SkillsStyle',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=4,
                spaceBefore=0,
                fontName='Helvetica',
                leftIndent=0,
                rightIndent=0,
                alignment=0,  # Left align
                leading=12
            )
            
            # Date style for cover letters - 10px (consistent with body text)
            date_style = ParagraphStyle(
                'DateStyle',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=12,
                spaceBefore=0,
                alignment=0,  # Left align
                fontName='Helvetica',
                textColor=black,
                leading=12
            )
            
            # Section-specific styles - ALL CONSISTENT AT 10px
            # Professional Summary style
            summary_style = ParagraphStyle(
                'SummaryStyle',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=4,
                spaceBefore=0,
                fontName='Helvetica',
                leftIndent=0,
                rightIndent=0,
                alignment=0,
                leading=12,
                textColor=black
            )
            
            # Work Experience style
            experience_style = ParagraphStyle(
                'ExperienceStyle',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=4,
                spaceBefore=0,
                fontName='Helvetica',
                leftIndent=0,
                rightIndent=0,
                alignment=0,
                leading=12,
                textColor=black
            )
            
            # Education style
            education_style = ParagraphStyle(
                'EducationStyle',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=4,
                spaceBefore=0,
                fontName='Helvetica',
                leftIndent=0,
                rightIndent=0,
                alignment=0,
                leading=12,
                textColor=black
            )
            
            # Projects style
            projects_style = ParagraphStyle(
                'ProjectsStyle',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=4,
                spaceBefore=0,
                fontName='Helvetica',
                leftIndent=0,
                rightIndent=0,
                alignment=0,
                leading=12,
                textColor=black
            )
            
            # Certifications style
            certifications_style = ParagraphStyle(
                'CertificationsStyle',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=4,
                spaceBefore=0,
                fontName='Helvetica',
                leftIndent=0,
                rightIndent=0,
                alignment=0,
                leading=12,
                textColor=black
            )
            
            # Build content with improved formatting
            story = []
            lines = content.split('\n')
            i = 0
            in_skills_section = False
            current_section = None
            
            while i < len(lines):
                line = lines[i].strip()
                if not line:
                    i += 1
                    continue
                
                # Check if it's a name (first non-empty line, all caps, no special chars)
                if i == 0 or (line.isupper() and len(line) > 3 and not any(char in line for char in ['|', '@', '.', ':', '•'])):
                    story.append(Paragraph(line, name_style))
                    story.append(Spacer(1, 6))  # Slightly more space after name
                
                # Check if it's contact info (contains email, phone, or LinkedIn)
                elif ('@' in line or 'linkedin.com' in line.lower() or 
                      (any(char.isdigit() for char in line) and len(line) > 10) or
                      'github.com' in line.lower()):
                    story.append(Paragraph(line, contact_style))
                    story.append(Spacer(1, 8))  # More space after contact info
                
                # Check if it's a section header
                elif (line.isupper() and len(line) > 3 and 
                      any(section in line for section in [
                          'PROFESSIONAL SUMMARY', 'WORK EXPERIENCE', 'EDUCATION', 
                          'SKILLS', 'PROJECTS', 'CERTIFICATIONS', 'EXTRACURRICULAR ACTIVITIES',
                          'AWARDS & ACHIEVEMENTS', 'VOLUNTEER', 'LANGUAGES', 'INTERESTS',
                          'TECHNICAL SKILLS', 'PROFESSIONAL EXPERIENCE'
                      ])):
                    in_skills_section = 'SKILLS' in line
                    
                    # Set current section for styling
                    if 'PROFESSIONAL SUMMARY' in line:
                        current_section = 'summary'
                    elif 'WORK EXPERIENCE' in line or 'PROFESSIONAL EXPERIENCE' in line:
                        current_section = 'experience'
                    elif 'EDUCATION' in line:
                        current_section = 'education'
                    elif 'PROJECTS' in line:
                        current_section = 'projects'
                    elif 'CERTIFICATIONS' in line:
                        current_section = 'certifications'
                    else:
                        current_section = 'other'
                    
                    story.append(Paragraph(line, section_style))
                    # Add line under section header (matching popup)
                    story.append(Paragraph('_' * 50, ParagraphStyle(
                        'LineStyle',
                        parent=styles['Normal'],
                        fontSize=8,
                        spaceAfter=4,  # Consistent 4px spacing
                        spaceBefore=0,
                        fontName='Helvetica',
                        alignment=0,  # Left align
                        textColor=black
                    )))
                    story.append(Spacer(1, 4))  # Consistent 4px spacing
                
                # Check if it's a date (for cover letters)
                elif (len(line.split()) <= 3 and 
                      any(month in line for month in ['January', 'February', 'March', 'April', 'May', 'June',
                                                      'July', 'August', 'September', 'October', 'November', 'December'])):
                    story.append(Paragraph(line, date_style))
                
                # Check if it's a job title/company line (contains |)
                elif '|' in line and len(line.split('|')) >= 2:
                    story.append(Paragraph(line, job_title_style))
                    story.append(Spacer(1, 4))  # Consistent 4px spacing
                
                # Check if it's a bullet point
                elif line.startswith('•') or line.startswith('-'):
                    story.append(Paragraph(line, bullet_style))
                    story.append(Spacer(1, 4))  # Consistent 4px spacing
                
                # Check if it's in skills section (format differently)
                elif in_skills_section and ':' in line:
                    story.append(Paragraph(line, skills_style))
                    story.append(Spacer(1, 4))  # Consistent 4px spacing
                
                # Regular text with section-specific styling
                else:
                    # Choose style based on current section
                    if current_section == 'summary':
                        story.append(Paragraph(line, summary_style))
                    elif current_section == 'experience':
                        story.append(Paragraph(line, experience_style))
                    elif current_section == 'education':
                        story.append(Paragraph(line, education_style))
                    elif current_section == 'projects':
                        story.append(Paragraph(line, projects_style))
                    elif current_section == 'certifications':
                        story.append(Paragraph(line, certifications_style))
                    else:
                        story.append(Paragraph(line, normal_style))
                    
                    story.append(Spacer(1, 4))  # Consistent 4px spacing
                
                i += 1
            
            # Build PDF
            doc.build(story)
            pdf_io.seek(0)
            return pdf_io.getvalue()
            
        except ImportError:
            raise HTTPException(status_code=500, detail="PDF generation requires reportlab. Please install it: pip install reportlab")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error creating PDF: {str(e)}")
    
    return b""

@router.post("/export")
async def export_document(request: ExportRequest):
    """Export document in specified format"""
    
    try:
        print(f"🔄 Backend received export request:")
        print(f"📄 Filename: {request.filename}")
        print(f"📋 Format: {request.format}")
        print(f"📝 Content preview: {request.content[:200]}...")
        print(f"📊 Content length: {len(request.content)}")
        
        # Validate file format
        if request.format not in ["docx", "pdf"]:
            raise HTTPException(status_code=400, detail="Unsupported file format. Supported formats: docx, pdf")
        
        # Create file content
        print(f"🚀 Creating {request.format.upper()} file...")
        file_content = create_download_file(request.content, request.filename, request.format)
        print(f"✅ File created successfully, size: {len(file_content)} bytes")
        
        if not file_content:
            raise HTTPException(status_code=500, detail="Failed to create file content")
        
        # Determine MIME type
        mime_types = {
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "pdf": "application/pdf"
        }
        
        mime_type = mime_types.get(request.format, "application/octet-stream")
        
        # Create streaming response
        def iter_content():
            yield file_content
        
        return StreamingResponse(
            iter_content(),
            media_type=mime_type,
            headers={
                "Content-Disposition": f"attachment; filename={request.filename}.{request.format}"
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error exporting document: {str(e)}")

@router.get("/formats")
async def get_supported_formats():
    """Get list of supported export formats"""
    return {
        "supported_formats": [
            {
                "format": "docx",
                "description": "Microsoft Word document",
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            },
            {
                "format": "pdf",
                "description": "Portable Document Format",
                "mime_type": "application/pdf"
            }
        ]
    }
