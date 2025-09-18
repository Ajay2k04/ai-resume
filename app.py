import streamlit as st
import pandas as pd
import os
import io
import re
from typing import Dict, List, Optional, Tuple
import json
from datetime import datetime, timedelta
import requests
from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
import uvicorn
import threading

# Import required libraries
try:
    from jobspy import scrape_jobs
    import pdfplumber
    from docx import Document
    from openai import AzureOpenAI
    from dotenv import load_dotenv
except ImportError as e:
    st.error(f"Missing required package: {e}")
    st.stop()

# Load environment variables
load_dotenv()

# FastAPI Models for Cover Letter Endpoint
class ContactInfo(BaseModel):
    email: str
    phone: str
    linkedin: Optional[str] = ""
    website: Optional[str] = ""

class Experience(BaseModel):
    title: str
    company: str
    duration: str
    description: str

class ResumeData(BaseModel):
    name: str
    contact_info: ContactInfo
    skills: List[str]
    experience: List[Experience]
    education: List[str]
    summary: str
    projects: List[str]

class CoverLetterRequest(BaseModel):
    resume_data: ResumeData
    job_description: str
    job_title: str
    company_name: str

class CoverLetterResponse(BaseModel):
    cover_letter: str
    status: str
    message: str

# Initialize FastAPI app for cover letter endpoint
cover_letter_app = FastAPI(
    title="Cover Letter Generation API",
    description="AI-powered cover letter generation service",
    version="1.0.0"
)

@cover_letter_app.get("/")
async def root():
    """Root endpoint"""
    return {
        "message": "Cover Letter Generation API",
        "version": "1.0.0",
        "status": "running"
    }

@cover_letter_app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat()
    }

@cover_letter_app.post("/generate-cover-letter", response_model=CoverLetterResponse)
async def generate_cover_letter_endpoint(request: CoverLetterRequest):
    """
    Generate a tailored cover letter based on resume data and job description
    """
    try:
        # Validate required environment variables
        required_vars = [
            'AZURE_OPENAI_API_KEY',
            'AZURE_OPENAI_ENDPOINT',
            'AZURE_OPENAI_DEPLOYMENT_NAME'
        ]
        
        missing_vars = [var for var in required_vars if not os.getenv(var)]
        if missing_vars:
            raise HTTPException(
                status_code=500, 
                detail=f"Missing environment variables: {', '.join(missing_vars)}"
            )
        
        # Convert Pydantic model to dict for the existing function
        resume_dict = request.resume_data.dict()
        
        # Generate cover letter using existing function
        # We need to pass the job title and company name to the function
        # Let's modify the call to include these parameters
        cover_letter = generate_ai_content_with_job_details(
            resume_dict,
            request.job_description,
            request.job_title,
            request.company_name,
            "cover_letter"
        )
        
        if not cover_letter:
            raise HTTPException(status_code=500, detail="Failed to generate cover letter")
        
        return CoverLetterResponse(
            cover_letter=cover_letter,
            status="success",
            message="Cover letter generated successfully"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")

# Function to start FastAPI server in background
def start_cover_letter_api():
    """Start the FastAPI server for cover letter generation"""
    uvicorn.run(cover_letter_app, host="0.0.0.0", port=8001, log_level="error")

# Start the API server in background if not already running
def ensure_api_running():
    """Ensure the cover letter API is running"""
    try:
        # Check if API is already running
        response = requests.get("http://localhost:8001/health", timeout=2)
        if response.status_code == 200:
            return True
    except:
        pass
    
    # Start API server in background thread
    try:
        api_thread = threading.Thread(target=start_cover_letter_api, daemon=True)
        api_thread.start()
        # Give it a moment to start
        import time
        time.sleep(2)
        return True
    except:
        return False

# Configure Streamlit page
st.set_page_config(
    page_title="QuantiPeak Recruitment Assistant",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Custom CSS for modern professional UI
st.markdown("""
<style>
    /* Import Google Fonts */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    
    /* Global Styles */
    * {
        font-family: 'Inter', sans-serif;
    }
    
    .main > div {
        padding-top: 0;
    }
    
    .stApp {
        max-width: 100%;
        margin: 0;
        background: #000000;
    }
    
    .main .block-container {
        background: transparent;
        padding: 0;
        margin: 0;
        max-width: 100%;
    }
    
    /* Header Section */
    .header-section {
        background: linear-gradient(135deg, #000000 0%, #1a1a1a 100%);
        color: white;
        padding: 3rem 2rem;
        text-align: center;
        margin-bottom: 2rem;
        border-radius: 0 0 20px 20px;
        border: 1px solid #333333;
    }
    
    .header-section h1 {
        font-size: 2.5rem;
        font-weight: 700;
        margin: 0;
        color: white;
    }
    
    .header-section p {
        font-size: 1.1rem;
        margin: 1rem 0 0 0;
        opacity: 0.9;
    }
    
    /* Step Progress */
    .step-progress {
        display: flex;
        justify-content: center;
        margin: 2rem 0;
        padding: 0 2rem;
    }
    
    .step-item {
        display: flex;
        align-items: center;
        margin: 0 1rem;
    }
    
    .step-circle {
        width: 40px;
        height: 40px;
        border-radius: 50%;
        display: flex;
        align-items: center;
        justify-content: center;
        font-weight: 600;
        margin-right: 0.5rem;
    }
    
    .step-active {
        background: #3b82f6;
        color: white;
    }
    
    .step-completed {
        background: #10b981;
        color: white;
    }
    
    .step-pending {
        background: #e5e7eb;
        color: #6b7280;
    }
    
    /* Main Content */
    .content-container {
        max-width: 1200px;
        margin: 0 auto;
        padding: 0 2rem;
    }
    
    .step-container {
        background: #1a1a1a;
        border-radius: 16px;
        padding: 2.5rem;
        margin: 2rem 0;
        box-shadow: 0 4px 6px -1px rgba(255, 255, 255, 0.1);
        border: 1px solid #333333;
    }
    
    .step-container h2 {
        color: #ffffff;
        font-size: 1.5rem;
        font-weight: 600;
        margin-bottom: 1.5rem;
        border-bottom: 2px solid #00ff88;
        padding-bottom: 0.5rem;
    }
    
    /* Job Cards */
    .job-card {
        background: #2a2a2a;
        border: 1px solid #444444;
        border-radius: 12px;
        padding: 1.5rem;
        margin: 1rem 0;
        transition: all 0.3s ease;
        cursor: pointer;
    }
    
    .job-card:hover {
        border-color: #00ff88;
        box-shadow: 0 10px 25px rgba(0, 255, 136, 0.2);
        transform: translateY(-2px);
    }
    
    .job-title {
        font-size: 1.2rem;
        font-weight: 600;
        color: #ffffff;
        margin-bottom: 0.5rem;
    }
    
    .job-company {
        color: #00ff88;
        font-weight: 500;
        margin-bottom: 0.5rem;
    }
    
    .job-location {
        color: #cccccc;
        font-size: 0.9rem;
        margin-bottom: 1rem;
    }
    
    .job-description {
        color: #aaaaaa;
        line-height: 1.6;
        margin-bottom: 1rem;
    }
    
    /* Buttons */
    .stButton > button {
        background: linear-gradient(135deg, #00ff88 0%, #00cc6a 100%);
        color: #000000;
        border: none;
        border-radius: 8px;
        padding: 0.75rem 2rem;
        font-weight: 600;
        font-size: 1rem;
        transition: all 0.3s ease;
        box-shadow: 0 4px 6px rgba(0, 255, 136, 0.3);
    }
    
    .stButton > button:hover {
        transform: translateY(-1px);
        box-shadow: 0 8px 15px rgba(0, 255, 136, 0.4);
        background: linear-gradient(135deg, #00ff88 0%, #00ff88 100%);
    }
    
    /* Form Elements */
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea,
    .stNumberInput > div > div > input,
    .stSelectbox > div > div {
        border: 2px solid #444444;
        border-radius: 8px;
        padding: 0.75rem;
        font-size: 1rem;
        transition: all 0.3s ease;
        background: #2a2a2a;
        color: #ffffff;
    }
    
    .stTextInput > div > div > input:focus,
    .stTextArea > div > div > textarea:focus,
    .stNumberInput > div > div > input:focus {
        border-color: #00ff88;
        box-shadow: 0 0 0 3px rgba(0, 255, 136, 0.2);
    }
    
    /* File Upload */
    .stFileUploader > div {
        border: 2px dashed #00ff88;
        border-radius: 12px;
        padding: 2rem;
        text-align: center;
        background: #1a1a1a;
        transition: all 0.3s ease;
    }
    
    .stFileUploader > div:hover {
        border-color: #00ff88;
        background: #2a2a2a;
    }
    
    /* Checkbox */
    .stCheckbox > div > div {
        background: #2a2a2a;
        border: 1px solid #444444;
        border-radius: 8px;
        padding: 1rem;
    }
    
    /* Success/Error Messages */
    .success-message {
        background: #1a2a1a;
        border: 1px solid #00ff88;
        color: #00ff88;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
    }
    
    .error-message {
        background: #2a1a1a;
        border: 1px solid #ff4444;
        color: #ff4444;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
    }
    
    /* Sidebar */
    .stSidebar {
        background: #1a1a1a;
    }
    
    /* Hide Streamlit branding */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    
    /* Custom spacing */
    .stMarkdown {
        margin-bottom: 1rem;
    }
    
    /* Download buttons */
    .download-section {
        background: #1a1a1a;
        border-radius: 12px;
        padding: 2rem;
        margin: 2rem 0;
        border: 1px solid #333333;
    }
    
    .download-section h3 {
        color: #ffffff;
        margin-bottom: 1rem;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'step' not in st.session_state:
    st.session_state.step = 1
if 'jobs_data' not in st.session_state:
    st.session_state.jobs_data = None
if 'resume_data' not in st.session_state:
    st.session_state.resume_data = None
if 'selected_job' not in st.session_state:
    st.session_state.selected_job = None
if 'generated_resume' not in st.session_state:
    st.session_state.generated_resume = None
if 'generated_cover_letter' not in st.session_state:
    st.session_state.generated_cover_letter = None
if 'show_cover_letter_form' not in st.session_state:
    st.session_state.show_cover_letter_form = False
if 'current_page' not in st.session_state:
    st.session_state.current_page = 'home'

def check_azure_config() -> bool:
    """Check if Azure OpenAI configuration is available"""
    required_vars = [
        'AZURE_OPENAI_API_KEY',
        'AZURE_OPENAI_ENDPOINT',
        'AZURE_OPENAI_DEPLOYMENT_NAME'
    ]
    
    missing_vars = [var for var in required_vars if not os.getenv(var)]
    
    if missing_vars:
        st.error(f"Missing environment variables: {', '.join(missing_vars)}")
        st.info("Please set the following environment variables:")
        for var in missing_vars:
            st.code(f"{var}=your_value_here")
        return False
    return True

def scrape_jobs_linkedin(
    job_titles: List[str],
    location: str,
    remote_only: bool,
    posted_within_days: int,
    num_results: int
) -> pd.DataFrame:
    """Scrape jobs from LinkedIn using jobspy with optimized performance"""
    try:
        import time
        
        # Create progress indicators
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        # Optimize search terms for better performance
        if job_titles:
            # Limit to 2 job titles for faster scraping
            search_terms = job_titles[:2]
            search_term = " OR ".join(search_terms)
        else:
            search_term = "Software Engineer"
        
        # Optimize parameters for speed
        optimized_results = min(num_results, 30)  # Cap at 30 for faster results
        optimized_days = min(posted_within_days, 14)  # Cap at 14 days for speed
        
        # Update progress
        progress_bar.progress(20)
        status_text.text("🔍 Initializing job search...")
        time.sleep(0.3)
        
        # Calculate date filter
        date_posted = datetime.now() - timedelta(days=optimized_days)
        
        progress_bar.progress(40)
        status_text.text("📡 Connecting to LinkedIn...")
        
        # Scrape jobs with optimized settings
        try:
            jobs = scrape_jobs(
                site_name=["linkedin"],
                search_term=search_term,
                location=location if location else "United States",
                results_wanted=optimized_results,
                hours_old=optimized_days * 24,
                country_indeed="us",
                linkedin_fetch_description=True
            )
        except Exception as scrape_error:
            st.warning(f"LinkedIn scraping encountered an issue: {str(scrape_error)}")
            # Try with even more conservative settings
            jobs = scrape_jobs(
                site_name=["linkedin"],
                search_term=search_term,
                location=location if location else "United States",
                results_wanted=min(optimized_results, 15),
                hours_old=optimized_days * 24,
                country_indeed="us",
                linkedin_fetch_description=False  # Skip descriptions for speed
            )
        
        progress_bar.progress(70)
        status_text.text("📊 Processing job data...")
        
        if jobs.empty:
            progress_bar.progress(100)
            status_text.text("❌ No jobs found")
            time.sleep(1)
            progress_bar.empty()
            status_text.empty()
            return pd.DataFrame()
        
        # Filter for remote jobs if requested
        if remote_only:
            jobs = jobs[jobs['location'].str.contains('remote|Remote|REMOTE', case=False, na=False)]
        
        # Select and rename columns
        job_columns = ['title', 'company', 'location', 'job_url', 'description', 'company_url']
        available_columns = [col for col in job_columns if col in jobs.columns]
        
        result_df = jobs[available_columns].copy()
        result_df.columns = ['Job Title', 'Company', 'Location', 'Job URL', 'Description', 'Company LinkedIn/Website']
        
        # Clean up the data efficiently
        result_df = result_df.fillna('N/A')
        result_df = result_df.drop_duplicates(subset=['Job URL'])
        
        # Filter out jobs without descriptions (if descriptions were fetched)
        if 'Description' in result_df.columns:
            result_df = result_df[result_df['Description'] != 'N/A']
            result_df = result_df[result_df['Description'].str.len() > 30]  # Reduced threshold for speed
        
        # Limit final results for better performance
        result_df = result_df.head(min(len(result_df), 25))
        
        progress_bar.progress(100)
        status_text.text(f"✅ Found {len(result_df)} jobs successfully!")
        time.sleep(1)
        
        # Clear progress indicators
        progress_bar.empty()
        status_text.empty()
        
        return result_df
        
    except Exception as e:
        st.error(f"Error scraping jobs: {str(e)}")
        # Clear any progress indicators
        try:
            progress_bar.empty()
            status_text.empty()
        except:
            pass
        return pd.DataFrame()

def parse_pdf_resume(file_content: bytes) -> Dict:
    """Parse PDF resume and extract structured data"""
    try:
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            text = ""
            for page in pdf.pages:
                text += page.extract_text() or ""
        
        return parse_resume_text(text)
    except Exception as e:
        st.error(f"Error parsing PDF: {str(e)}")
        return {}

def parse_docx_resume(file_content: bytes) -> Dict:
    """Parse DOCX resume and extract structured data"""
    try:
        doc = Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return parse_resume_text(text)
    except Exception as e:
        st.error(f"Error parsing DOCX: {str(e)}")
        return {}

def parse_resume_text(text: str) -> Dict:
    """Parse resume text and extract structured information"""
    resume_data = {
        'name': '',
        'contact_info': {
            'email': '',
            'phone': '',
            'linkedin': '',
            'website': ''
        },
        'skills': [],
        'experience': [],
        'education': [],
        'summary': '',
        'projects': []
    }
    
    # Extract name (usually first line or after "Name:")
    lines = text.split('\n')
    for line in lines[:5]:  # Check first 5 lines
        line = line.strip()
        if line and not any(keyword in line.lower() for keyword in ['email', 'phone', 'linkedin', 'github', 'portfolio']):
            resume_data['name'] = line
            break
    
    # Extract email
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = re.findall(email_pattern, text)
    if emails:
        resume_data['contact_info']['email'] = emails[0]
    
    # Extract phone
    phone_pattern = r'(\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})'
    phones = re.findall(phone_pattern, text)
    if phones:
        resume_data['contact_info']['phone'] = ''.join(phones[0])
    
    # Extract LinkedIn
    linkedin_pattern = r'linkedin\.com/in/[\w-]+'
    linkedin_matches = re.findall(linkedin_pattern, text, re.IGNORECASE)
    if linkedin_matches:
        resume_data['contact_info']['linkedin'] = f"https://{linkedin_matches[0]}"
    
    # Extract website/portfolio
    website_pattern = r'(?:https?://)?(?:www\.)?[a-zA-Z0-9-]+\.[a-zA-Z]{2,}(?:/[^\s]*)?'
    websites = re.findall(website_pattern, text)
    for website in websites:
        if 'linkedin' not in website.lower() and 'github' not in website.lower():
            resume_data['contact_info']['website'] = website
            break
    
    # Enhanced skills extraction
    skill_keywords = [
        'python', 'java', 'javascript', 'typescript', 'react', 'angular', 'vue', 'node.js', 'sql',
        'machine learning', 'ai', 'artificial intelligence', 'data science', 'analytics', 'aws', 'azure', 'gcp',
        'docker', 'kubernetes', 'git', 'agile', 'scrum', 'project management', 'html', 'css', 'bootstrap',
        'mongodb', 'postgresql', 'mysql', 'redis', 'elasticsearch', 'tensorflow', 'pytorch', 'pandas',
        'numpy', 'scikit-learn', 'flask', 'django', 'fastapi', 'spring', 'express', 'rest api', 'graphql',
        'microservices', 'ci/cd', 'jenkins', 'terraform', 'ansible', 'linux', 'unix', 'bash', 'powershell'
    ]
    
    text_lower = text.lower()
    for skill in skill_keywords:
        if skill in text_lower:
            resume_data['skills'].append(skill.title())
    
    # Extract experience with duration parsing
    experience_patterns = [
        r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s+at\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)',
        r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s*-\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)',
        r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s*@\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)',
        r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s*,\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)',
        r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s*\|([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)'
    ]
    
    # Look for duration patterns
    duration_patterns = [
        r'(\d+)\s*(?:years?|yrs?)\s*(?:\d+\s*months?)?',
        r'(\d+)\s*months?',
        r'(\d+)\s*-\s*(\d+)\s*years?',
        r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s*\d{4}\s*-\s*(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s*\d{4}',
        r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s*\d{4}\s*-\s*Present'
    ]
    
    for pattern in experience_patterns:
        experiences = re.findall(pattern, text)
        for exp in experiences[:3]:  # Limit to 3 experiences
            if len(exp[0]) > 3 and len(exp[1]) > 2:  # Filter out short matches
                # Clean up the extracted data
                title = exp[0].strip()
                company = exp[1].strip()
                
                # Skip if it looks like location or other non-company data
                if not any(word in company.lower() for word in ['bangalore', 'mumbai', 'delhi', 'hyderabad', 'chennai', 'pune', 'kolkata', 'salem', 'tamil']):
                    # Try to find duration in the text around this experience
                    duration = 'Duration not specified'
                    for dur_pattern in duration_patterns:
                        duration_match = re.search(dur_pattern, text, re.IGNORECASE)
                        if duration_match:
                            duration = duration_match.group(0)
                            break
                    
                    resume_data['experience'].append({
                        'title': title,
                        'company': company,
                        'duration': duration,
                        'description': f'Contributed to software development projects using modern technologies and best practices.'
                    })
    
    # If no experience found, try to extract from project descriptions
    if not resume_data['experience']:
        project_lines = [line for line in lines if any(keyword in line.lower() for keyword in ['developed', 'built', 'created', 'implemented'])]
        if project_lines:
            resume_data['experience'].append({
                'title': 'Software Developer',
                'company': 'Various Projects',
                'duration': 'Project-based experience',
                'description': 'Developed multiple software applications and projects using various technologies.'
            })
    
    # Extract education with better patterns
    education_patterns = [
        r'(Bachelor|Master|PhD|B\.S\.|M\.S\.|Ph\.D\.|B\.A\.|M\.A\.)\s+[A-Za-z\s]+',
        r'[A-Za-z]+\s+University',
        r'[A-Za-z]+\s+College',
        r'[A-Za-z]+\s+Institute'
    ]
    
    for pattern in education_patterns:
        educations = re.findall(pattern, text)
        for edu in educations[:2]:  # Limit to 2 education entries
            resume_data['education'].append(edu)
    
    # Extract projects with better formatting
    project_keywords = ['project', 'developed', 'built', 'created', 'implemented', 'designed']
    lines = text.split('\n')
    for i, line in enumerate(lines):
        line_lower = line.lower()
        if any(keyword in line_lower for keyword in project_keywords):
            if len(line.strip()) > 20:  # Filter out short lines
                # Clean up the project description
                clean_line = line.strip()
                # Add proper spacing between words
                clean_line = re.sub(r'([a-z])([A-Z])', r'\1 \2', clean_line)
                clean_line = re.sub(r'([a-z])(\d)', r'\1 \2', clean_line)
                clean_line = re.sub(r'(\d)([A-Z])', r'\1 \2', clean_line)
                resume_data['projects'].append(clean_line)
    
    # Create a basic summary
    if resume_data['skills']:
        skills_text = ', '.join(resume_data['skills'][:5])
        resume_data['summary'] = f"Experienced professional with expertise in {skills_text}. Passionate about technology and innovation with a strong foundation in software development and problem-solving."
    
    return resume_data

def call_cover_letter_api(resume_data: Dict, job_description: str, job_title: str, company_name: str) -> str:
    """Call the FastAPI cover letter generation endpoint"""
    try:
        url = "http://localhost:8001/generate-cover-letter"
        
        # Convert resume data to the format expected by the API
        api_resume_data = {
            "name": resume_data.get('name', ''),
            "contact_info": {
                "email": resume_data.get('contact_info', {}).get('email', ''),
                "phone": resume_data.get('contact_info', {}).get('phone', ''),
                "linkedin": resume_data.get('contact_info', {}).get('linkedin', ''),
                "website": resume_data.get('contact_info', {}).get('website', '')
            },
            "skills": resume_data.get('skills', []),
            "experience": resume_data.get('experience', []),
            "education": resume_data.get('education', []),
            "summary": resume_data.get('summary', ''),
            "projects": resume_data.get('projects', [])
        }
        
        payload = {
            "resume_data": api_resume_data,
            "job_description": job_description,
            "job_title": job_title,
            "company_name": company_name
        }
        
        response = requests.post(url, json=payload, timeout=30)
        response.raise_for_status()
        
        result = response.json()
        return result.get('cover_letter', '')
        
    except requests.exceptions.ConnectionError:
        st.error("❌ Could not connect to the cover letter API. Please ensure the API server is running on http://localhost:8001")
        return ""
    except requests.exceptions.HTTPError as e:
        st.error(f"❌ API Error: {e}")
        return ""
    except Exception as e:
        st.error(f"❌ Error calling API: {str(e)}")
        return ""

def generate_ai_content_with_job_details(resume_data: Dict, job_description: str, job_title: str, company_name: str, content_type: str) -> str:
    """Generate tailored resume or cover letter using Azure OpenAI with specific job details"""
    try:
        client = AzureOpenAI(
            api_key=os.getenv('AZURE_OPENAI_API_KEY'),
            api_version=os.getenv('AZURE_OPENAI_API_VERSION', '2023-07-01-preview'),
            azure_endpoint=os.getenv('AZURE_OPENAI_ENDPOINT')
        )
        
        # Get candidate details
        candidate_name = resume_data.get('name', 'Candidate')
        candidate_email = resume_data.get('contact_info', {}).get('email', 'candidate@email.com')
        candidate_phone = resume_data.get('contact_info', {}).get('phone', 'Phone Number')
        candidate_linkedin = resume_data.get('contact_info', {}).get('linkedin', 'LinkedIn Profile')
        
        if content_type == "resume":
            # Clean the resume data for better formatting
            resume_data_cleaned = json.dumps(resume_data, indent=2)
            
            prompt = f"""
            You are a professional resume writer. 
            Generate a clean, ATS-friendly resume for {candidate_name} applying for the role of {job_title} at {company_name}.

            CANDIDATE INFO (from resume):
            {resume_data_cleaned}

            JOB DESCRIPTION:
            {job_description}

            STRICT RULES:
            - Use ONLY candidate's actual information. No fake companies, roles, or achievements.
            - If a section has no data, OMIT it completely.
            - Use only explicitly listed skills, tools, or experiences. Do not invent.
            - Tailor wording to match {job_title} at {company_name} by aligning with the job description keywords.
            - Do NOT include location or address.
            - CRITICAL: Do NOT use asterisk (*) symbols anywhere in the resume - use only bullet points (•).
            - Do NOT use "|" symbols in project names or anywhere in the resume.
            - Output only the final resume. No extra notes, explanations, or markdown.

            RESUME FORMAT:

            {candidate_name.upper()}
            {candidate_email} | {candidate_phone} | {candidate_linkedin}

            PROFESSIONAL SUMMARY
            2–3 sentences summarizing the candidate's actual skills/experience. 
            Tailored specifically to {job_title} at {company_name}. 
            Do not use generic claims like "experienced" unless supported by the data.

            TECHNICAL SKILLS
            Programming Languages: [from resume data]
            Frameworks & Libraries: [from resume data]
            Tools & Technologies: [from resume data]
            Databases: [from resume data]

            PROFESSIONAL EXPERIENCE
            [For each real experience found:]
            Job Title | Company | Duration (exactly as in resume data)
            • Responsibility or achievement from resume
            • Responsibility or achievement from resume
            • Responsibility or achievement from resume

            EDUCATION
            [List only if present in resume data, format: Degree | Institution | Year]

            PROJECTS
            [For each project in resume data:]
            Project Name
            • Description of what was built/accomplished
            • Key feature or measurable outcome

            FORMATTING REQUIREMENTS:
            - Left align all text.
            - Use consistent bullet points (•) ONLY - NEVER use asterisk (*) symbols.
            - Maintain clear spacing between sections.
            - Ensure ATS compatibility with standard section headers.
            - Do not include "IMPORTANT NOTES" or commentary.
            - CRITICAL: Do NOT use asterisk (*) symbols anywhere in the resume - use only bullet points (•).
            - Do NOT use "|" symbols in project names or anywhere in the resume.
            """
        else:  # cover letter
            from datetime import datetime
            current_date = datetime.now().strftime("%B %d, %Y")
            
            # Clean the resume data for better formatting
            resume_data_cleaned = json.dumps(resume_data, indent=2)
            
            prompt = f"""
            You are a professional career coach and resume writer. 
            Generate a cover letter for {candidate_name}, applying for the role of {job_title} at {company_name}.

            CANDIDATE INFO (from resume):
            {resume_data_cleaned}

            JOB DESCRIPTION:
            {job_description}

            STRICT RULES:
            - Use ONLY candidate's real information. No fake experience, projects, or achievements.
            - If a section of information is missing, skip it completely.
            - Do NOT invent company values or mission statements. Use only details in the JD.
            - Be specific: use keywords from the JD that match the candidate's actual skills.
            - Do not repeat the resume content verbatim; rewrite it as a persuasive narrative.
            - Output ONLY the cover letter in plain text, no markdown, no extra notes.

            FORMAT:
            {current_date}

            {company_name}
            Human Resources Department

            Dear Hiring Manager,

            [Intro: Tailored opening using candidate's real skills & background]

            [Body 1: Highlight candidate's skills matching job requirements]

            [Body 2: Explain motivation to join {company_name}, using details from JD]

            [Body 3: Explain what candidate can contribute based on real achievements]

            Sincerely,
            {candidate_name}

            Contact:
            Email: {candidate_email} | Phone: {candidate_phone} | LinkedIn: {candidate_linkedin}
            """
        
        response = client.chat.completions.create(
            model=os.getenv('AZURE_OPENAI_DEPLOYMENT_NAME'),
            messages=[
                {"role": "system", "content": "You are an expert career coach and professional resume writer. You specialize in creating compelling, tailored resumes and cover letters that get candidates hired. CRITICAL: Use ONLY the candidate's actual information provided. Do NOT create fake companies, fake experience, fake achievements, or fake details. If information is missing, simply omit that section. Focus on tailoring the existing real information to match the job requirements. NEVER use asterisk (*) symbols in resumes - use only bullet points (•)."},
                
                {"role": "user", "content": prompt}
            ],
            max_tokens=4000,
            temperature=0.2
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        print(f"Error generating {content_type}: {str(e)}")
        return ""

def generate_ai_content(resume_data: Dict, job_description: str, content_type: str) -> str:
    """Generate tailored resume or cover letter using Azure OpenAI"""
    try:
        client = AzureOpenAI(
            api_key=os.getenv('AZURE_OPENAI_API_KEY'),
            api_version=os.getenv('AZURE_OPENAI_API_VERSION', '2023-07-01-preview'),
            azure_endpoint=os.getenv('AZURE_OPENAI_ENDPOINT')
        )
        
        # Extract job details from the job description
        job_title = "Software Engineer"  # Default fallback
        company_name = "the company"  # Default fallback
        
        # Try to extract job title and company from the job description
        if st.session_state.selected_job:
            job_title = st.session_state.selected_job.get('Job Title', 'Software Engineer')
            company_name = st.session_state.selected_job.get('Company', 'the company')
        
        # Get candidate details
        candidate_name = resume_data.get('name', 'Candidate')
        candidate_email = resume_data.get('contact_info', {}).get('email', 'candidate@email.com')
        candidate_phone = resume_data.get('contact_info', {}).get('phone', 'Phone Number')
        candidate_linkedin = resume_data.get('contact_info', {}).get('linkedin', 'LinkedIn Profile')
        
        if content_type == "resume":
            prompt = f"""
            Create a professional, well-formatted resume for {candidate_name} applying for the {job_title} position at {company_name}. 

            CANDIDATE INFORMATION:
            {json.dumps(resume_data, indent=2)}

            TARGET JOB DESCRIPTION:
            {job_description}

            CRITICAL REQUIREMENTS:
            1. Use ONLY the candidate's actual information from the resume data provided
            2. Do NOT create fake companies, fake experience, or fake achievements
            3. Do NOT include location in the resume
            4. Create a clean, professional format with proper spacing
            5. Use bullet points for better readability
            6. Tailor the content to match the job requirements
            7. Use keywords from the job description
            8. Make it ATS-friendly with clear section headers
            9. Do NOT include any formatting notes or explanations at the end
            10. Do NOT include any "IMPORTANT FORMATTING" or "NOTES" sections

            FORMAT THE RESUME EXACTLY LIKE THIS WITH PROPER ALIGNMENT:

            {candidate_name.upper()}
            {candidate_email} | {candidate_phone} | {candidate_linkedin}

            PROFESSIONAL SUMMARY
            Write 2-3 sentences based on the candidate's actual skills and experience from their resume. Do NOT use generic words like "experienced" unless they actually have significant experience. Use their real skills and tailor to the {job_title} role at {company_name}. Be specific about their actual capabilities.

            TECHNICAL SKILLS
            Programming Languages: List actual programming languages from resume data
            Frameworks & Libraries: List actual frameworks from resume data
            Tools & Technologies: List actual tools from resume data
            Databases: List actual databases from resume data

            PROFESSIONAL EXPERIENCE
            Only include experience that was actually found in the resume data. Format each entry as:
            Job Title | Company Name | Actual Duration from resume - do NOT use "2+ years" unless that's what's in the resume
            - Description based on actual skills and projects from their resume
            - Another relevant point from their actual experience
            - Third relevant point from their actual work

            If no proper experience found, skip this section entirely

            EDUCATION
            Only include education that was actually found in the resume data. If no education was found, skip this section

            PROJECTS
            [Only include projects that were actually found in the resume data. Format as:]
            [Project Name] | [Technologies Used]
            • [Description of what was accomplished]
            • [Key features or achievements]

            [If no projects found, skip this section entirely]

            PROJECT FORMATTING RULES:
            - Clean up project descriptions by adding proper spacing
            - Separate words that are run together (e.g., "Developedresponsive" → "Developed responsive")
            - Use proper punctuation and capitalization
            - Make descriptions clear and professional

            FORMATTING REQUIREMENTS:
            - Use proper spacing between sections
            - Left-align all content
            - Use consistent bullet points (•)
            - Ensure clean, professional formatting
            - Make it ATS-friendly with clear section headers
            - Do NOT include any formatting notes or explanations at the end
            - Do NOT include any "IMPORTANT FORMATTING" or "NOTES" sections

            """
        else:  # cover letter
            from datetime import datetime
            current_date = datetime.now().strftime("%B %d, %Y")
            
            prompt = f"""
            Write a professional cover letter for {candidate_name} applying for the {job_title} position at {company_name}.

            CANDIDATE INFORMATION:
            {json.dumps(resume_data, indent=2)}

            TARGET JOB DESCRIPTION:
            {job_description}

            CRITICAL REQUIREMENTS:
            1. Use ONLY the candidate's actual information from the resume data
            2. Do NOT create fake experience, fake projects, or fake achievements
            3. Use the candidate's actual name: {candidate_name}
            4. Use the exact job title: {job_title}
            5. Use the exact company name: {company_name}
            6. Show enthusiasm based on their real skills and experience
            7. NO fake details, NO made-up experience, NO fictional achievements

            FORMAT THE COVER LETTER EXACTLY LIKE THIS:

            {current_date}

            {company_name}
            Human Resources Department

            Dear Hiring Manager,

            I am writing to express my interest in the {job_title} position at {company_name}. [Write opening based on candidate's actual skills and experience from their resume. Do NOT use generic words like "experienced" unless they actually have significant experience. Be specific about their real capabilities.]

            [Body paragraph 1: Highlight the candidate's actual skills that match the job requirements. Use only skills found in their resume data. Tailor the language to match the job description keywords and requirements.]

            [Body paragraph 2: Show knowledge of the company and role, and explain why the candidate wants to work there based on their real interests and skills. Use specific details about the company and role from the job description.]

            [Body paragraph 3: Emphasize what the candidate can contribute based on their actual experience and skills. Use their real project descriptions and achievements from their resume.]

            I am excited about the opportunity to contribute to {company_name} and would welcome the chance to discuss how my background and skills align with your needs. Thank you for considering my application.

            Sincerely,
            {candidate_name}

            ---
            Contact Information:
            Email: {candidate_email}
            Phone: {candidate_phone}
            LinkedIn: {candidate_linkedin}

            IMPORTANT: 
            - Use only real information from the candidate's resume
            - Do not create fake experience or achievements
            - Focus on their actual skills and how they match the job requirements
            - Be honest about their background and experience level
            """
        
        response = client.chat.completions.create(
            model=os.getenv('AZURE_OPENAI_DEPLOYMENT_NAME'),
            messages=[
                {"role": "system", "content": "You are an expert career coach and professional resume writer. You specialize in creating compelling, tailored resumes and cover letters that get candidates hired. CRITICAL: Use ONLY the candidate's actual information provided. Do NOT create fake companies, fake experience, fake achievements, or fake details. If information is missing, simply omit that section. Focus on tailoring the existing real information to match the job requirements. NEVER use asterisk (*) symbols in resumes - use only bullet points (•)."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=4000,
            temperature=0.2
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        st.error(f"Error generating {content_type}: {str(e)}")
        return ""

def create_download_file(content: str, filename: str, file_type: str) -> bytes:
    """Create downloadable file content"""
    if file_type == "txt":
        return content.encode('utf-8')
    elif file_type == "docx":
        # Create a simple DOCX file
        doc = Document()
        doc.add_paragraph(content)
        
        # Save to bytes
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io.getvalue()
    elif file_type == "pdf":
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib.colors import black, darkblue
            
            # Create PDF in memory
            pdf_io = io.BytesIO()
            doc = SimpleDocTemplate(pdf_io, pagesize=letter, 
                                  leftMargin=0.75*inch, rightMargin=0.75*inch,
                                  topMargin=0.75*inch, bottomMargin=0.75*inch)
            
            # Create custom styles
            styles = getSampleStyleSheet()
            
            # Name style
            name_style = ParagraphStyle(
                'NameStyle',
                parent=styles['Heading1'],
                fontSize=14,
                spaceAfter=6,
                alignment=0,  # Left align
                textColor=black,
                fontName='Helvetica-Bold'
            )
            
            # Contact style
            contact_style = ParagraphStyle(
                'ContactStyle',
                parent=styles['Normal'],
                fontSize=9,
                spaceAfter=12,
                alignment=0,  # Left align
                fontName='Helvetica'
            )
            
            # Section header style
            section_style = ParagraphStyle(
                'SectionStyle',
                parent=styles['Heading2'],
                fontSize=10,
                spaceAfter=6,
                spaceBefore=10,
                textColor=black,
                fontName='Helvetica-Bold',
                alignment=0,  # Left align
                leftIndent=0,
                rightIndent=0
            )
            
            # Normal text style
            normal_style = ParagraphStyle(
                'NormalStyle',
                parent=styles['Normal'],
                fontSize=9,
                spaceAfter=2,
                fontName='Helvetica',
                leftIndent=0,
                rightIndent=0,
                alignment=0  # Left align
            )
            
            # Bullet point style
            bullet_style = ParagraphStyle(
                'BulletStyle',
                parent=styles['Normal'],
                fontSize=9,
                spaceAfter=2,
                fontName='Helvetica',
                leftIndent=15,
                rightIndent=0,
                alignment=0  # Left align
            )
            
            # Build content with proper formatting
            story = []
            lines = content.split('\n')
            i = 0
            
            while i < len(lines):
                line = lines[i].strip()
                if not line:
                    i += 1
                    continue
                
                # Check if it's a name (first non-empty line, all caps, no special chars)
                if i == 0 or (line.isupper() and len(line) > 3 and not any(char in line for char in ['|', '@', '.'])):
                    story.append(Paragraph(line, name_style))
                    story.append(Spacer(1, 8))
                
                # Check if it's contact info (contains email or phone)
                elif '@' in line or (any(char.isdigit() for char in line) and len(line) > 10):
                    story.append(Paragraph(line, contact_style))
                    story.append(Spacer(1, 12))
                
                # Check if it's a section header (all caps, common resume sections)
                elif (line.isupper() and len(line) > 3 and 
                      any(section in line for section in ['PROFESSIONAL SUMMARY', 'TECHNICAL SKILLS', 
                                                         'PROFESSIONAL EXPERIENCE', 'EDUCATION', 'PROJECTS'])):
                    story.append(Paragraph(line, section_style))
                    story.append(Spacer(1, 8))
                
                # Check if it's a date (for cover letters)
                elif (len(line.split()) <= 3 and 
                      any(month in line for month in ['January', 'February', 'March', 'April', 'May', 'June',
                                                      'July', 'August', 'September', 'October', 'November', 'December'])):
                    date_style = ParagraphStyle(
                        'DateStyle',
                        parent=styles['Normal'],
                        fontSize=9,
                        spaceAfter=10,
                        alignment=0,  # Left align
                        fontName='Helvetica',
                        textColor=black
                    )
                    story.append(Paragraph(line, date_style))
                
                # Check if it's contact info section (for cover letters)
                elif line.startswith('---') or 'Contact Information:' in line:
                    contact_header_style = ParagraphStyle(
                        'ContactHeaderStyle',
                        parent=styles['Normal'],
                        fontSize=10,
                        spaceAfter=5,
                        spaceBefore=10,
                        fontName='Helvetica-Bold'
                    )
                    story.append(Paragraph(line, contact_header_style))
                
                # Check if it's a bullet point
                elif line.startswith('•') or line.startswith('-'):
                    story.append(Paragraph(line, bullet_style))
                
                # Check if it's a job title/company line (contains |)
                elif '|' in line and len(line.split('|')) >= 2:
                    story.append(Paragraph(line, normal_style))
                    story.append(Spacer(1, 4))
                
                # Regular text
                else:
                    story.append(Paragraph(line, normal_style))
                
                i += 1
            
            # Build PDF
            doc.build(story)
            pdf_io.seek(0)
            return pdf_io.getvalue()
            
        except ImportError:
            st.error("PDF generation requires reportlab. Please install it: pip install reportlab")
            return b""
        except Exception as e:
            st.error(f"Error creating PDF: {str(e)}")
            return b""
    
    return b""

def render_navigation():
    """Render navigation bar"""
    st.markdown("""
    <div style="background: #1a1a1a; padding: 1rem; border-radius: 8px; margin-bottom: 2rem; border: 1px solid #333333;">
        <div style="display: flex; justify-content: space-between; align-items: center;">
            <div style="color: #00ff88; font-size: 1.5rem; font-weight: bold;">🚀 QuantiPeak</div>
            <div style="display: flex; gap: 1rem;">
                <a href="?page=home" style="color: #ffffff; text-decoration: none; padding: 0.5rem 1rem; border-radius: 4px; background: #2a2a2a;">🏠 Home</a>
                <a href="?page=cover-letter" style="color: #ffffff; text-decoration: none; padding: 0.5rem 1rem; border-radius: 4px; background: #2a2a2a;">📝 Cover Letter</a>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

def render_cover_letter_page():
    """Render the cover letter page"""
    st.markdown("## 📝 Cover Letter Generator")
    st.markdown("Generate professional, AI-powered cover letters tailored to specific job positions.")
    
    # Resume upload section
    if not st.session_state.resume_data:
        st.markdown("### 📄 Step 1: Upload Your Resume")
        uploaded_file = st.file_uploader(
            "Choose a resume file",
            type=['pdf', 'docx'],
            help="Upload your resume in PDF or DOCX format"
        )
        
        if uploaded_file is not None:
            if st.button("📄 Parse Resume", type="primary"):
                with st.spinner("Parsing resume..."):
                    file_content = uploaded_file.read()
                    
                    if uploaded_file.type == "application/pdf":
                        resume_data = parse_pdf_resume(file_content)
                    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        resume_data = parse_docx_resume(file_content)
                    else:
                        st.error("Unsupported file type")
                        resume_data = {}
                    
                    if resume_data:
                        st.session_state.resume_data = resume_data
                        st.success("✅ Resume parsed successfully!")
                        st.rerun()
    else:
        st.success("✅ Resume loaded successfully!")
        
        # Job details input
        st.markdown("### 💼 Step 2: Job Details")
        col1, col2 = st.columns(2)
        
        with col1:
            job_title = st.text_input("Job Title", placeholder="e.g., Software Engineer")
        with col2:
            company_name = st.text_input("Company Name", placeholder="e.g., Google, Microsoft")
        
        job_description = st.text_area(
            "Job Description", 
            height=200,
            placeholder="Paste the complete job description here..."
        )
        
        # Generate button
        if st.button("🎯 Generate Cover Letter", type="primary", use_container_width=True):
            if job_title and company_name and job_description:
                with st.spinner("🤖 AI is generating your personalized cover letter..."):
                    cover_letter = call_cover_letter_api(
                        st.session_state.resume_data,
                        job_description,
                        job_title,
                        company_name
                    )
                    if cover_letter:
                        st.session_state.generated_cover_letter = cover_letter
                        st.success("✅ Cover letter generated successfully!")
                    else:
                        st.error("❌ Failed to generate cover letter")
            else:
                st.error("⚠️ Please fill in all fields: Job Title, Company Name, and Job Description")
        
        # Display and download generated cover letter
        if st.session_state.generated_cover_letter:
            st.markdown("### 📄 Generated Cover Letter")
            st.markdown(f'<div class="generated-content">{st.session_state.generated_cover_letter}</div>', unsafe_allow_html=True)
            
            st.markdown("### 💾 Download Options")
            col1, col2, col3 = st.columns(3)
            
            with col1:
                cover_txt = create_download_file(st.session_state.generated_cover_letter, "cover_letter", "txt")
                st.download_button("📄 Download TXT", cover_txt, "cover_letter.txt", "text/plain", use_container_width=True)
            
            with col2:
                cover_docx = create_download_file(st.session_state.generated_cover_letter, "cover_letter", "docx")
                st.download_button("📝 Download DOCX", cover_docx, "cover_letter.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            
            with col3:
                cover_pdf = create_download_file(st.session_state.generated_cover_letter, "cover_letter", "pdf")
                if cover_pdf:
                    st.download_button("📋 Download PDF", cover_pdf, "cover_letter.pdf", "application/pdf", use_container_width=True)
        
        # Reset button
        if st.button("🔄 Start Over", type="secondary", use_container_width=True):
            for key in ['resume_data', 'generated_cover_letter']:
                st.session_state[key] = None
            st.rerun()

def render_home_page():
    """Render the home page with full recruitment workflow"""
    st.markdown("""
    <div class="header-section">
        <h1>🚀 QuantiPeak Recruitment Assistant</h1>
        <p>Streamline your recruitment workflow with AI-powered job matching and document generation</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Check Azure configuration
    if not check_azure_config():
        st.stop()
    
    # Show full recruitment workflow directly
    render_full_workflow()

def render_full_workflow():
    """Render the full recruitment workflow"""
    # Step Progress Indicator
    steps = ["Scrape Jobs", "Upload Resume", "Select Job", "Generate Documents", "Download"]
    current_step = st.session_state.step
    
    # Create step progress using Streamlit columns
    cols = st.columns(5)
    for i, (col, step) in enumerate(zip(cols, steps), 1):
        with col:
            if i < current_step:
                st.markdown(f"""
                <div style="text-align: center; padding: 10px;">
                    <div style="width: 40px; height: 40px; border-radius: 50%; background: #00ff88; color: #000000; display: flex; align-items: center; justify-content: center; margin: 0 auto 5px; font-weight: bold;">✓</div>
                    <div style="font-size: 12px; color: #00ff88; font-weight: 500;">{step}</div>
                </div>
                """, unsafe_allow_html=True)
            elif i == current_step:
                st.markdown(f"""
                <div style="text-align: center; padding: 10px;">
                    <div style="width: 40px; height: 40px; border-radius: 50%; background: #00ff88; color: #000000; display: flex; align-items: center; justify-content: center; margin: 0 auto 5px; font-weight: bold;">{i}</div>
                    <div style="font-size: 12px; color: #00ff88; font-weight: 500;">{step}</div>
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown(f"""
                <div style="text-align: center; padding: 10px;">
                    <div style="width: 40px; height: 40px; border-radius: 50%; background: #444444; color: #888888; display: flex; align-items: center; justify-content: center; margin: 0 auto 5px; font-weight: bold;">{i}</div>
                    <div style="font-size: 12px; color: #888888; font-weight: 500;">{step}</div>
                </div>
                """, unsafe_allow_html=True)
    
    # Step 1: Job Scraping
    if st.session_state.step == 1:
        st.markdown("## 🔍 Step 1: Search LinkedIn Jobs")
        
        col1, col2 = st.columns(2)
        
        with col1:
            job_titles_text = st.text_area(
                "Job Titles (one per line)",
                value="Software Engineer\nData Scientist\nProduct Manager",
                height=100
            )
            job_titles = [title.strip() for title in job_titles_text.split('\n') if title.strip()]
            
            location = st.text_input("Location", value="United States")
            remote_only = st.checkbox("Remote Only")
        
        with col2:
            posted_within_days = st.number_input("Posted within (days)", min_value=1, max_value=30, value=7)
            num_results = st.number_input("Number of results", min_value=1, max_value=100, value=20)
        
        if st.button("🔍 Search Jobs", type="primary"):
            # Show optimization tips
            with st.expander("💡 Performance Tips", expanded=False):
                st.markdown("""
                **For faster results:**
                - Use 1-2 specific job titles
                - Limit results to 20-30 jobs
                - Search within last 7-14 days
                - Use broader locations (e.g., "United States" instead of specific cities)
                """)
            
            # Add timeout warning
            st.info("⏱️ Job scraping may take 30-60 seconds. Please be patient...")
            
            jobs_df = scrape_jobs_linkedin(
                job_titles, location, remote_only, posted_within_days, num_results
            )
            
            if not jobs_df.empty:
                st.session_state.jobs_data = jobs_df
                st.session_state.step = 2
                st.success(f"✅ Found {len(jobs_df)} jobs! Proceeding to Step 2.")
                st.rerun()
            else:
                st.error("No jobs found. Please try different search criteria.")
                st.markdown("""
                **Try these alternatives:**
                - Use more general job titles (e.g., "Software Engineer" instead of "Senior Full Stack React Developer")
                - Increase the time range (try 14-30 days)
                - Use broader locations
                - Reduce the number of results requested
                """)
        
        if st.session_state.jobs_data is not None:
            st.markdown("### 📋 Searched Jobs Preview")
            st.dataframe(st.session_state.jobs_data.head(), width='stretch')
    
    # Step 2: Resume Upload
    elif st.session_state.step == 2:
        st.markdown("## 📄 Step 2: Upload and Parse Resume")
        
        uploaded_file = st.file_uploader(
            "Choose a resume file",
            type=['pdf', 'docx'],
            help="Upload your resume in PDF or DOCX format"
        )
        
        if uploaded_file is not None:
            if st.button("📄 Parse Resume", type="primary"):
                with st.spinner("Parsing resume..."):
                    file_content = uploaded_file.read()
                    
                    if uploaded_file.type == "application/pdf":
                        resume_data = parse_pdf_resume(file_content)
                    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        resume_data = parse_docx_resume(file_content)
                    else:
                        st.error("Unsupported file type")
                        resume_data = {}
                    
                    if resume_data:
                        st.session_state.resume_data = resume_data
                        st.session_state.step = 3
                        st.success("✅ Resume parsed successfully! Proceeding to Step 3.")
                        st.rerun()
        
        if st.session_state.resume_data is not None:
            st.markdown("### 📋 Parsed Resume Data")
            st.json(st.session_state.resume_data)
    
    # Step 3: Job Selection
    elif st.session_state.step == 3:
        st.markdown("## 🎯 Step 3: Select a Job Posting")
        
        if st.session_state.jobs_data is not None:
            jobs_df = st.session_state.jobs_data
            
            # Create job selection interface
            for idx, job in jobs_df.iterrows():
                job_html = f'''
                <div class="job-card">
                    <div class="job-title">{job['Job Title']}</div>
                    <div class="job-company">{job['Company']}</div>
                    <div class="job-location">{job['Location']}</div>
                    <div class="job-description">{job['Description'][:200]}...</div>
                </div>
                '''
                st.markdown(job_html, unsafe_allow_html=True)
                
                col1, col2, col3 = st.columns([1, 1, 1])
                with col2:
                    if st.button(f"Select This Job", key=f"select_{idx}", type="primary"):
                        st.session_state.selected_job = job.to_dict()
                        st.session_state.step = 4
                        st.success("✅ Job selected! Proceeding to Step 4.")
                        st.rerun()
        else:
            st.error("No jobs data available. Please go back to Step 1.")
    
    # Step 4: Generate Documents
    elif st.session_state.step == 4:
        st.markdown("## ✨ Step 4: Generate Tailored Resume & Cover Letter")
        
        if st.session_state.selected_job and st.session_state.resume_data:
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button("📝 Generate Tailored Resume", type="primary"):
                    with st.spinner("Generating tailored resume..."):
                        tailored_resume = generate_ai_content(
                            st.session_state.resume_data,
                            st.session_state.selected_job['Description'],
                            "resume"
                        )
                        if tailored_resume:
                            st.session_state.generated_resume = tailored_resume
                            st.success("✅ Tailored resume generated!")
            
            with col2:
                if st.button("💌 Generate Cover Letter", type="primary"):
                    with st.spinner("Generating cover letter..."):
                        cover_letter = generate_ai_content(
                            st.session_state.resume_data,
                            st.session_state.selected_job['Description'],
                            "cover_letter"
                        )
                        if cover_letter:
                            st.session_state.generated_cover_letter = cover_letter
                            st.success("✅ Cover letter generated!")
            
            # Display generated content
            if st.session_state.generated_resume:
                st.markdown("### 📝 Generated Tailored Resume")
                st.text_area("Resume", value=st.session_state.generated_resume, height=400)
            
            if st.session_state.generated_cover_letter:
                st.markdown("### 💌 Generated Cover Letter")
                st.text_area("Cover Letter", value=st.session_state.generated_cover_letter, height=400)
            
            # Proceed to download step
            if st.session_state.generated_resume and st.session_state.generated_cover_letter:
                if st.button("📥 Proceed to Download", type="primary"):
                    st.session_state.step = 5
                    st.rerun()
        else:
            st.error("Missing job selection or resume data. Please complete previous steps.")
    
    # Step 5: Download
    elif st.session_state.step == 5:
        st.markdown("## 📥 Step 5: Download Generated Documents")
        
        if st.session_state.generated_resume and st.session_state.generated_cover_letter:
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("### 📄 Download Tailored Resume")
                resume_txt = create_download_file(st.session_state.generated_resume, "tailored_resume", "txt")
                resume_docx = create_download_file(st.session_state.generated_resume, "tailored_resume", "docx")
                resume_pdf = create_download_file(st.session_state.generated_resume, "tailored_resume", "pdf")
                
                st.download_button(
                    label="📄 Download as TXT",
                    data=resume_txt,
                    file_name="tailored_resume.txt",
                    mime="text/plain"
                )
                
                st.download_button(
                    label="📄 Download as DOCX",
                    data=resume_docx,
                    file_name="tailored_resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                st.download_button(
                    label="📄 Download as PDF",
                    data=resume_pdf,
                    file_name="tailored_resume.pdf",
                    mime="application/pdf"
                )
            
            with col2:
                st.markdown("### 💌 Download Cover Letter")
                cover_txt = create_download_file(st.session_state.generated_cover_letter, "cover_letter", "txt")
                cover_docx = create_download_file(st.session_state.generated_cover_letter, "cover_letter", "docx")
                cover_pdf = create_download_file(st.session_state.generated_cover_letter, "cover_letter", "pdf")
                
                st.download_button(
                    label="💌 Download as TXT",
                    data=cover_txt,
                    file_name="cover_letter.txt",
                    mime="text/plain"
                )
                
                st.download_button(
                    label="💌 Download as DOCX",
                    data=cover_docx,
                    file_name="cover_letter.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                st.download_button(
                    label="💌 Download as PDF",
                    data=cover_pdf,
                    file_name="cover_letter.pdf",
                    mime="application/pdf"
                )
            
            # Reset button
            if st.button("🔄 Start New Process", type="secondary"):
                # Reset session state
                for key in ['step', 'jobs_data', 'resume_data', 'selected_job', 'generated_resume', 'generated_cover_letter']:
                    if key == 'step':
                        st.session_state[key] = 1
                    else:
                        st.session_state[key] = None
                st.rerun()
        else:
            st.error("No generated documents available. Please complete Step 4.")
    
    # Navigation buttons
    st.markdown("---")
    col1, col2, col3 = st.columns([1, 1, 1])
    
    with col1:
        if st.session_state.step > 1:
            if st.button("← Previous Step", use_container_width=True):
                st.session_state.step -= 1
                st.rerun()
    
    with col2:
        # Removed Back to Home button
        pass
    
    with col3:
        if st.session_state.step < 5:
            if st.button("Next Step →", use_container_width=True):
                st.session_state.step += 1
                st.rerun()
    
    # Sidebar with current job info
    if st.session_state.selected_job:
        with st.sidebar:
            st.markdown("## 🎯 Selected Job")
            job = st.session_state.selected_job
            st.write(f"**{job['Job Title']}**")
            st.write(f"*{job['Company']}*")
            st.write(f"📍 {job['Location']}")
            st.write(f"🔗 [View Job]({job['Job URL']})")

def main():
    """Main application function"""
    # Check for page parameter in URL
    query_params = st.query_params
    page = query_params.get('page', 'home')
    
    # Update current page
    st.session_state.current_page = page
    
    # Render navigation
    render_navigation()
    
    # Route to appropriate page
    if page == 'cover-letter':
        render_cover_letter_page()
    else:
        render_home_page()

if __name__ == "__main__":
    main()
